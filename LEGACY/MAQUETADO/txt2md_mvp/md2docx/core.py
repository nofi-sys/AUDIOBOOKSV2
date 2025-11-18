#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md2docx_core.py â€” MVP
--------------------
Convierte un archivo Markdown (.md) a .docx con control de maquetado (tamaÃ±o de pÃ¡gina, mÃ¡rgenes y estilos tipogrÃ¡ficos por tipo de pÃ¡rrafo).
Pensado para KDP (Amazon). Arquitectura modular para importarlo como mÃ³dulo desde otra app.
"""
import re
import json
import copy
from pathlib import Path
from dataclasses import dataclass
from typing import List, Dict, Tuple, Any, Optional, Set

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -----------------------------
# Modelos de datos y defaults
# -----------------------------

STYLE_KEYS_ORDER = [
    "heading1",
    "heading2",
    "heading3",
    "subtitle",
    "paragraph",
    "blockquote",
    "codeblock",
    "bullet",
    "numbered",
]

DEFAULT_PAGE = {
    "preset": "KDP 5x8 in",
    "width_in": 5.0,
    "height_in": 8.0,
    "margin_top_in": 0.75,
    "margin_bottom_in": 0.75,
    "margin_left_in": 0.75,
    "margin_right_in": 0.75,
}

DEFAULT_STYLESET: Dict[str, Dict[str, Any]] = {
    "heading1": {
        "font_name": "Garamond",
        "font_size_pt": 16,
        "align": "justify",
        "bold": True,
        "italic": False,
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
        "page_break_before": True,
    },
    "heading2": {
        "font_name": "Garamond",
        "font_size_pt": 14,
        "align": "justify",
        "bold": True,
        "italic": False,
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
        "page_break_before": True,
    },
    "heading3": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "justify",
        "bold": True,
        "italic": False,
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "subtitle": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "center",
        "bold": False,
        "italic": True,
        "line_spacing": 1.0,
        "space_before_pt": 0,
        "space_after_pt": 12,
        "first_line_indent_in": 0.0,
    },
    "paragraph": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "justify",
        "bold": False,
        "italic": False,
        "line_spacing": 1.0,  # simple
        "space_before_pt": 0,
        "space_after_pt": 6,  # requisito: espacio despuÃ©s de pÃ¡rrafo por defecto
        "first_line_indent_in": 0.0,
    },
    "blockquote": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "justify",
        "bold": False,
        "italic": False,
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
        "left_indent_in": 0.5
    },
    "poem": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "center",
        "bold": False,
        "italic": True,
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "letter": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "left",
        "bold": False,
        "italic": False,
        "line_spacing": 1.0,
        "space_before_pt": 0,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "codeblock": {
        "font_name": "Courier New",
        "font_size_pt": 10,
        "align": "justify",
        "bold": False,
        "italic": False,
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "bullet": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "justify",
        "bold": False,
        "italic": False,
        "line_spacing": 1.0,
        "space_before_pt": 0,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "numbered": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "justify",
        "bold": False,
        "italic": False,
        "line_spacing": 1.0,
        "space_before_pt": 0,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "_global": {
        "generate_toc": False,
        "page_number_alignment": "center",
        "page_number_font_name": None,
        "page_number_font_size_pt": None,
        "header_left_text": "",
        "header_right_text": "",
        "header_font_name": None,
        "header_font_size_pt": None,
        "header_italic": True,
        "header_distance_in": None,
        "footer_distance_in": None,
        "header_space_after_pt": 12,
        "footer_space_before_pt": 12,
        "header_alignment": "mirror",
        "chapter_padding_top_lines": 4,
        "chapter_padding_bottom_lines": 1,
    },
}

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# -----------------------------
# Parser sencillo de Markdown
# -----------------------------

@dataclass
class Block:
    type: str  # 'heading1','heading2','heading3','paragraph','blockquote','codeblock','bullet','numbered','hr'
    content: Any  # str, list[str] (para listas), etc.

INLINE_TOKEN_RE = re.compile(
    r"(`[^`]+`)|(\*\*[^*]+\*\*)|(__[^_]+__)|(\*[^*]+\*)|(_[^_]+_)|(~~[^~]+~~)|(\+\+[^+]+\+\+)"
)

def tokenize_inline(text: str) -> List[Tuple[str, str]]:
    tokens = []
    pos = 0
    for m in INLINE_TOKEN_RE.finditer(text):
        start, end = m.span()
        if start > pos:
            tokens.append(("text", text[pos:start]))
        group = m.group(0)
        if group.startswith("`"):
            tokens.append(("code", group[1:-1]))
        elif group.startswith("**") and group.endswith("**"):
            tokens.append(("bold", group[2:-2]))
        elif group.startswith("__") and group.endswith("__"):
            tokens.append(("bold", group[2:-2]))
        elif group.startswith("*") and group.endswith("*"):
            tokens.append(("italic", group[1:-1]))
        elif group.startswith("_") and group.endswith("_"):
            tokens.append(("italic", group[1:-1]))
        elif group.startswith("~~") and group.endswith("~~"):
            tokens.append(("strike", group[2:-2]))
        elif group.startswith("++") and group.endswith("++"):
            tokens.append(("underline", group[2:-2]))
        pos = end
    if pos < len(text):
        tokens.append(("text", text[pos:]))
    return tokens

FENCE_RE = re.compile(r"^```(\w+)?\s*$")
ATX_H_RE = re.compile(r"^(#{1,6})\s+(.*)$")
HR_RE = re.compile(r"^(\*\s*\*\s*\*|-{3,}|_{3,})\s*$")
BULLET_RE = re.compile(r"^([\-\*\+])\s+(.*)$")
ORDERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")

def parse_markdown_blocks(md_text: str) -> Tuple[List[Block], Set[str]]:
    all_lines = md_text.splitlines()
    i = 0

    # 1. Omitir el frontmatter de YAML si existe
    if all_lines and all_lines[0].strip() == '---':
        try:
            # Encontrar el '---' de cierre y empezar despuÃ©s
            i = all_lines.index('---', 1) + 1
        except ValueError:
            pass  # No hay '---' de cierre, procesar como si no hubiera frontmatter

    # 2. Omitir lÃ­neas en blanco y el tÃ­tulo del TOC ("Contenidos")
    while i < len(all_lines) and not all_lines[i].strip():
        i += 1
    if i < len(all_lines) and all_lines[i].strip().lower() in ('contenidos', 'contents'):
        i += 1
    while i < len(all_lines) and not all_lines[i].strip():
        i += 1

    # 3. Omitir los elementos de la lista del TOC generado por Markdown
    toc_item_re = re.compile(r'^\s*[\*\-\â€¢]\s+\[.+?\]\(#.+?\)\s*$')
    while i < len(all_lines) and toc_item_re.match(all_lines[i].strip()):
        i += 1

    # El resto del contenido se procesa desde el Ã­ndice 'i'
    lines = all_lines
    blocks: List[Block] = []
    used_styles: Set[str] = set()
    while i < len(lines):
        line = lines[i]

        m = FENCE_RE.match(line.strip())
        if m:
            lang = m.group(1) or ""
            i += 1
            buf = []
            while i < len(lines) and not FENCE_RE.match(lines[i].strip()):
                buf.append(lines[i])
                i += 1
            if i < len(lines) and FENCE_RE.match(lines[i].strip()):
                i += 1
            code_text = "\n".join(buf)
            blocks.append(Block("codeblock", code_text))
            used_styles.add("codeblock")
            continue

        if HR_RE.match(line.strip()):
            blocks.append(Block("hr", ""))
            i += 1
            continue

        m = ATX_H_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                blocks.append(Block("heading1", text))
                used_styles.add("heading1")
            elif level == 2:
                blocks.append(Block("heading2", text))
                used_styles.add("heading2")
            elif level == 3:
                blocks.append(Block("heading3", text))
                used_styles.add("heading3")
            else:
                blocks.append(Block("paragraph", text))
                used_styles.add("paragraph")
            i += 1
            continue

        m = BLOCKQUOTE_RE.match(line)
        if m:
            buf = [m.group(1)]
            i += 1
            while i < len(lines):
                m2 = BLOCKQUOTE_RE.match(lines[i])
                if m2:
                    buf.append(m2.group(1))
                    i += 1
                else:
                    break
            text = "\n".join(buf).strip()
            blocks.append(Block("blockquote", text))
            used_styles.add("blockquote")
            continue

        m = BULLET_RE.match(line)
        if m:
            items = [m.group(2)]
            i += 1
            while i < len(lines):
                m2 = BULLET_RE.match(lines[i])
                if m2:
                    items.append(m2.group(2))
                    i += 1
                else:
                    if lines[i].strip() == "":
                        i += 1
                        continue
                    break
            blocks.append(Block("bullet", items))
            used_styles.add("bullet")
            continue

        m = ORDERED_RE.match(line)
        if m:
            items = [m.group(2)]
            i += 1
            while i < len(lines):
                m2 = ORDERED_RE.match(lines[i])
                if m2:
                    items.append(m2.group(2))
                    i += 1
                else:
                    if lines[i].strip() == "":
                        i += 1
                        continue
                    break
            blocks.append(Block("numbered", items))
            used_styles.add("numbered")
            continue

        if line.strip() == "":
            i += 1
            continue
        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip() != "":
            buf.append(lines[i])
            i += 1
        text = "\n".join(buf).strip()
        blocks.append(Block("paragraph", text))
        used_styles.add("paragraph")

    if not used_styles:
        used_styles.add("paragraph")
    return blocks, used_styles

# -----------------------------
# ConstrucciÃ³n del DOCX
# -----------------------------

def apply_paragraph_format(p, style_cfg: Dict[str, Any]):
    pf = p.paragraph_format
    pf.alignment = ALIGN_MAP.get(style_cfg.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    ls = style_cfg.get("line_spacing", 1.0)
    if isinstance(ls, (int, float)):
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = ls
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    before = style_cfg.get("space_before_pt", 0)
    after = style_cfg.get("space_after_pt", 0)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    first_indent_in = style_cfg.get("first_line_indent_in", 0.0)
    left_indent_in = style_cfg.get("left_indent_in", 0.0)
    try:
        pf.first_line_indent = Inches(float(first_indent_in))
    except Exception:
        pf.first_line_indent = Inches(0.0)
    try:
        pf.left_indent = Inches(float(left_indent_in))
    except Exception:
        pf.left_indent = Inches(0.0)

def apply_run_format(run, make_bold=False, make_italic=False, make_underline=False, strike=False, font_name=None, font_size_pt=None):
    rfont = run.font
    rfont.bold = True if make_bold else None
    rfont.italic = True if make_italic else None
    rfont.underline = True if make_underline else None
    rfont.strike = True if strike else None
    if font_name:
        rfont.name = font_name
    if font_size_pt:
        rfont.size = Pt(font_size_pt)
    rfont.color.rgb = RGBColor(0,0,0)

def write_inline_runs(p, text: str, base_font_name: str, base_font_size: int, base_bold: bool = False, base_italic: bool = False):
    tokens = tokenize_inline(text)
    for kind, val in tokens:
        run = p.add_run()
        is_bold = base_bold or (kind == "bold")
        is_italic = base_italic or (kind == "italic")

        if kind == "text":
            run.text = val
            apply_run_format(run, make_bold=is_bold, make_italic=is_italic, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "bold":
            run.text = val
            apply_run_format(run, make_bold=True, make_italic=is_italic, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "italic":
            run.text = val
            apply_run_format(run, make_bold=is_bold, make_italic=True, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "underline":
            run.text = val
            apply_run_format(run, make_bold=is_bold, make_underline=True, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "strike":
            run.text = val
            apply_run_format(run, make_bold=is_bold, strike=True, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "code":
            run.text = val
            apply_run_format(run, font_name="Courier New", font_size_pt=max(9, int(base_font_size*0.9)))

def setup_page(document: Document, page_cfg: Dict[str, Any]):
    for section in document.sections:
        section.page_width = Inches(float(page_cfg["width_in"]))
        section.page_height = Inches(float(page_cfg["height_in"]))
        section.top_margin = Inches(float(page_cfg["margin_top_in"]))
        section.bottom_margin = Inches(float(page_cfg["margin_bottom_in"]))
        section.left_margin = Inches(float(page_cfg["margin_left_in"]))
        section.right_margin = Inches(float(page_cfg["margin_right_in"]))

def apply_style_to_named_style(document: Document, style_name: str, cfg: Dict[str, Any]):
    try:
        style = document.styles[style_name]
    except KeyError:
        style = document.styles['Normal']
    if hasattr(style, "font"):
        style.font.name = cfg.get("font_name", "Garamond")
        style.font.size = Pt(cfg.get("font_size_pt", 12))
        style.font.bold = cfg.get("bold", False)
        style.font.italic = cfg.get("italic", False)
        style.font.color.rgb = RGBColor(0,0,0)
    if hasattr(style, "paragraph_format"):
        pf = style.paragraph_format
        pf.alignment = ALIGN_MAP.get(cfg.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        pf.space_before = Pt(cfg.get("space_before_pt", 0))
        pf.space_after = Pt(cfg.get("space_after_pt", 6))
        pf.first_line_indent = Inches(cfg.get("first_line_indent_in", 0.0))
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = cfg.get("line_spacing", 1.0)

def build_document(blocks: List[Block], style_cfgs: Dict[str, Dict[str, Any]], page_cfg: Dict[str, Any]) -> Document:
    document = Document()
    setup_page(document, page_cfg)

    # AÃ±adir TOC si estÃ¡ configurado
    if style_cfgs.get("_global", {}).get("generate_toc", False):
        add_table_of_contents(document)

    apply_style_to_named_style(document, 'Normal', style_cfgs.get("paragraph", DEFAULT_STYLESET["paragraph"]))
    apply_style_to_named_style(document, 'Heading 1', style_cfgs.get("heading1", DEFAULT_STYLESET["heading1"]))
    apply_style_to_named_style(document, 'Heading 2', style_cfgs.get("heading2", DEFAULT_STYLESET["heading2"]))
    apply_style_to_named_style(document, 'Heading 3', style_cfgs.get("heading3", DEFAULT_STYLESET["heading3"]))
    apply_style_to_named_style(document, 'Quote', style_cfgs.get("blockquote", DEFAULT_STYLESET["blockquote"]))
    apply_style_to_named_style(document, 'List Bullet', style_cfgs.get("bullet", DEFAULT_STYLESET["bullet"]))
    apply_style_to_named_style(document, 'List Number', style_cfgs.get("numbered", DEFAULT_STYLESET["numbered"]))

    global_cfg = style_cfgs.get("_global", {})

    for i, blk in enumerate(blocks):
        if blk.type.startswith("heading"):
            level = blk.type.replace("heading", "")
            style_name = f"Heading {level}"
            cfg = style_cfgs.get(blk.type, DEFAULT_STYLESET.get(blk.type, {}))

            # AÃ±adir salto de pÃ¡gina si estÃ¡ configurado en los estilos
            if cfg.get("page_break_before", False) and i > 0:
                section = document.add_section(WD_SECTION_START.ODD_PAGE)
                # La nueva secciÃ³n hereda los mÃ¡rgenes y tamaÃ±o de la anterior,
                # pero es buena prÃ¡ctica re-aplicarlos si hubiera secciones con distintos tamaÃ±os.
                section.page_width = Inches(float(page_cfg["width_in"]))
                section.page_height = Inches(float(page_cfg["height_in"]))
                section.top_margin = Inches(float(page_cfg["margin_top_in"]))
                section.bottom_margin = Inches(float(page_cfg["margin_bottom_in"]))
                section.left_margin = Inches(float(page_cfg["margin_left_in"]))
                section.right_margin = Inches(float(page_cfg["margin_right_in"]))

            chapter_top_padding = 0
            chapter_bottom_padding = 0
            cfg_local = cfg
            if cfg.get("page_break_before", False):
                chapter_top_padding = _coerce_to_int(global_cfg.get("chapter_padding_top_lines")) or 0
                chapter_bottom_padding = _coerce_to_int(global_cfg.get("chapter_padding_bottom_lines")) or 0
                cfg_local = dict(cfg)
                cfg_local["space_before_pt"] = 0
                cfg_local["space_after_pt"] = 0
                for _ in range(chapter_top_padding):
                    blank = document.add_paragraph(style='Normal')
                    fmt = blank.paragraph_format
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.keep_with_next = True

            p = document.add_paragraph(style=style_name)
            apply_paragraph_format(p, cfg_local)
            write_inline_runs(p, blk.content, cfg_local["font_name"], cfg_local["font_size_pt"], base_bold=cfg_local.get("bold", False), base_italic=cfg_local.get("italic", False))
            p.paragraph_format.keep_with_next = True

            if chapter_bottom_padding:
                for _ in range(chapter_bottom_padding):
                    blank_after = document.add_paragraph(style='Normal')
                    fmt = blank_after.paragraph_format
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.keep_with_next = True
        elif blk.type == "paragraph":
            cfg = style_cfgs.get("paragraph", DEFAULT_STYLESET["paragraph"])
            p = document.add_paragraph(style='Normal')
            apply_paragraph_format(p, cfg)
            write_inline_runs(p, blk.content, cfg["font_name"], cfg["font_size_pt"], base_bold=cfg.get("bold", False), base_italic=cfg.get("italic", False))
        elif blk.type == "blockquote":
            cfg = style_cfgs.get("blockquote", DEFAULT_STYLESET["blockquote"])
            p = document.add_paragraph(style='Quote')
            apply_paragraph_format(p, cfg)
            write_inline_runs(p, blk.content, cfg["font_name"], cfg["font_size_pt"], base_bold=cfg.get("bold", False), base_italic=cfg.get("italic", False))
        elif blk.type == "codeblock":
            cfg = style_cfgs.get("codeblock", DEFAULT_STYLESET["codeblock"])
            p = document.add_paragraph(style='Normal')
            apply_paragraph_format(p, cfg)
            for line in blk.content.splitlines() or [""]:
                run = p.add_run(line)
                apply_run_format(run, font_name=cfg["font_name"], font_size_pt=cfg["font_size_pt"], make_bold=cfg.get("bold", False))
                p.add_run("\n")
        elif blk.type == "poem":
            cfg = style_cfgs.get("poem", DEFAULT_STYLESET["poem"])
            p = document.add_paragraph(style='Normal')
            apply_paragraph_format(p, cfg)
            lines = blk.content.splitlines() if isinstance(blk.content, str) else [str(blk.content)]
            if not lines:
                lines = [""]
            for i, line in enumerate(lines):
                if line:
                    write_inline_runs(p, line, cfg["font_name"], cfg["font_size_pt"], base_bold=cfg.get("bold", False), base_italic=cfg.get("italic", True))
                else:
                    run = p.add_run("")
                    apply_run_format(run, font_name=cfg["font_name"], font_size_pt=cfg["font_size_pt"], make_bold=cfg.get("bold", False), make_italic=cfg.get("italic", True))
                if i != len(lines) - 1:
                    p.add_run().add_break()
        elif blk.type == "letter":
            cfg = style_cfgs.get("letter", DEFAULT_STYLESET["letter"])
            paragraphs = re.split(r"\n\s*\n", blk.content.strip()) if isinstance(blk.content, str) else [str(blk.content)]
            if not paragraphs:
                paragraphs = [""]
            for para in paragraphs:
                p = document.add_paragraph(style='Normal')
                apply_paragraph_format(p, cfg)
                text = para.strip()
                if text:
                    write_inline_runs(p, text, cfg["font_name"], cfg["font_size_pt"], base_bold=cfg.get("bold", False), base_italic=cfg.get("italic", False))
                else:
                    run = p.add_run("")
                    apply_run_format(run, font_name=cfg["font_name"], font_size_pt=cfg["font_size_pt"], make_bold=cfg.get("bold", False))
        elif blk.type == "bullet":
            cfg = style_cfgs.get("bullet", DEFAULT_STYLESET["bullet"])
            for item in blk.content:
                p = document.add_paragraph(style='List Bullet')
                apply_paragraph_format(p, cfg)
                write_inline_runs(p, item, cfg["font_name"], cfg["font_size_pt"], base_bold=cfg.get("bold", False), base_italic=cfg.get("italic", False))
        elif blk.type == "numbered":
            cfg = style_cfgs.get("numbered", DEFAULT_STYLESET["numbered"])
            for item in blk.content:
                p = document.add_paragraph(style='List Number')
                apply_paragraph_format(p, cfg)
                write_inline_runs(p, item, cfg["font_name"], cfg["font_size_pt"], base_bold=cfg.get("bold", False), base_italic=cfg.get("italic", False))
        elif blk.type == "hr":
            p = document.add_paragraph()
            run = p.add_run("â€¢ â€¢ â€¢")
            apply_run_format(run, font_name='Garamond', font_size_pt=12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return document

def _clear_container_paragraphs(container) -> None:
    for paragraph in list(container.paragraphs):
        element = paragraph._p
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

def _append_page_field(run) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r"PAGE \* MERGEFORMAT"
    run._r.append(instr)

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_page_numbers(document: Document, style_cfgs: Dict[str, Dict[str, Any]]):
    """
    AÃ±ade numeraciÃ³n de pÃ¡gina al pie de todas las secciones utilizando la
    configuraciÃ³n global disponible en los estilos.
    """
    global_cfg = style_cfgs.get("_global", {})
    paragraph_cfg = style_cfgs.get("paragraph", DEFAULT_STYLESET["paragraph"])

    align_key = (global_cfg.get("page_number_alignment") or paragraph_cfg.get("align", "center")).lower()
    align = ALIGN_MAP.get(align_key, WD_ALIGN_PARAGRAPH.CENTER)
    font_name = global_cfg.get("page_number_font_name") or paragraph_cfg.get("font_name")
    font_size = _coerce_to_float(global_cfg.get("page_number_font_size_pt"))
    if font_size is None:
        font_size = _coerce_to_float(paragraph_cfg.get("font_size_pt")) or 12
    footer_space_before = _coerce_to_float(global_cfg.get("footer_space_before_pt")) or 0

    for section in document.sections:
        footer = section.footer
        even_footer = section.even_page_footer
        first_footer = section.first_page_footer

        # Primer pÃ¡gina en blanco
        first_footer.is_linked_to_previous = False
        _clear_container_paragraphs(first_footer)
        blank_first_footer = first_footer.add_paragraph()
        blank_first_footer.add_run("")
        first_fmt = blank_first_footer.paragraph_format
        first_fmt.space_before = Pt(0)
        first_fmt.space_after = Pt(0)

        for target in (footer, even_footer):
            target.is_linked_to_previous = False
            _clear_container_paragraphs(target)

            paragraph = target.add_paragraph()
            paragraph.alignment = align
            if footer_space_before:
                paragraph.paragraph_format.space_before = Pt(footer_space_before)

            run = paragraph.add_run()
            apply_run_format(run, font_name=font_name, font_size_pt=font_size)
            _append_page_field(run)

    ensure_update_fields(document)

def ensure_update_fields(document: Document) -> None:
    settings = document.settings
    element = settings.element
    update = element.find(qn('w:updateFields'))
    if update is None:
        update = OxmlElement('w:updateFields')
        update.set(qn('w:val'), 'true')
        element.append(update)
    else:
        update.set(qn('w:val'), 'true')

def _coerce_to_float(value: Any) -> Optional[float]:
    if value is None:
        return None
    try:
        text = str(value).strip()
    except Exception:
        return None
    if not text:
        return None
    if len(text) > 32:
        return None
    try:
        return float(text)
    except ValueError:
        return None

def _apply_global_defaults(style_cfgs: Dict[str, Dict[str, Any]], title: Optional[str] = None, author: Optional[str] = None) -> None:
    paragraph_cfg = style_cfgs.get("paragraph", DEFAULT_STYLESET["paragraph"])
    global_cfg = style_cfgs.setdefault("_global", {})

    if title and not global_cfg.get("header_left_text"):
        global_cfg["header_left_text"] = title
    if author and not global_cfg.get("header_right_text"):
        global_cfg["header_right_text"] = author

    global_cfg.setdefault("header_font_name", paragraph_cfg.get("font_name"))
    try:
        para_size = float(paragraph_cfg.get("font_size_pt", 12))
    except (TypeError, ValueError):
        para_size = 12.0
    global_cfg.setdefault("header_font_size_pt", max(para_size - 1, 6))
    if global_cfg.get("header_font_size_pt") is None:
        global_cfg["header_font_size_pt"] = max(para_size - 1, 6)

    global_cfg.setdefault("header_italic", True)
    global_cfg.setdefault("header_distance_in", DEFAULT_STYLESET["_global"]["header_distance_in"])
    global_cfg.setdefault("footer_distance_in", DEFAULT_STYLESET["_global"]["footer_distance_in"])
    global_cfg.setdefault("header_space_after_pt", DEFAULT_STYLESET["_global"]["header_space_after_pt"])
    global_cfg.setdefault("footer_space_before_pt", DEFAULT_STYLESET["_global"]["footer_space_before_pt"])
    global_cfg.setdefault("header_alignment", DEFAULT_STYLESET["_global"]["header_alignment"])
    global_cfg.setdefault("chapter_padding_top_lines", DEFAULT_STYLESET["_global"]["chapter_padding_top_lines"])
    global_cfg.setdefault("chapter_padding_bottom_lines", DEFAULT_STYLESET["_global"]["chapter_padding_bottom_lines"])

def _coerce_to_int(value: Any) -> Optional[int]:
    result = _coerce_to_float(value)
    if result is None:
        return None
    try:
        return max(0, int(round(result)))
    except (ValueError, OverflowError):
        return None

def configure_headers(document: Document, style_cfgs: Dict[str, Dict[str, Any]]):
    """
    Configura encabezados pares/impares segÃºn la configuraciÃ³n global y deja en
    blanco la primera pÃ¡gina de cada secciÃ³n (capÃ­tulos y pÃ¡ginas vacÃ­as).
    """
    global_cfg = style_cfgs.get("_global", {})
    paragraph_cfg = style_cfgs.get("paragraph", DEFAULT_STYLESET["paragraph"])

    left_text = (global_cfg.get("header_left_text") or "").strip()
    right_text = (global_cfg.get("header_right_text") or "").strip()
    font_name = global_cfg.get("header_font_name") or paragraph_cfg.get("font_name")
    font_size = _coerce_to_float(global_cfg.get("header_font_size_pt"))
    if font_size is None:
        font_size = _coerce_to_float(paragraph_cfg.get("font_size_pt")) or 12
    header_italic = bool(global_cfg.get("header_italic", False))
    header_space_after = _coerce_to_float(global_cfg.get("header_space_after_pt")) or 0
    header_alignment_setting = (global_cfg.get("header_alignment") or "mirror").lower()
    if header_alignment_setting == "left":
        odd_alignment = even_alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif header_alignment_setting == "center":
        odd_alignment = even_alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif header_alignment_setting == "right":
        odd_alignment = even_alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        odd_alignment = WD_ALIGN_PARAGRAPH.RIGHT
        even_alignment = WD_ALIGN_PARAGRAPH.LEFT

    header_distance = _coerce_to_float(global_cfg.get("header_distance_in"))
    footer_distance = _coerce_to_float(global_cfg.get("footer_distance_in"))

    any_header = bool(left_text or right_text)
    document.settings.odd_and_even_pages_header_footer = True

    for idx, section in enumerate(document.sections):
        section.different_first_page_header_footer = True
        try:
            if header_distance is not None:
                section.header_distance = Inches(header_distance)
        except Exception:
            pass
        try:
            if footer_distance is not None:
                section.footer_distance = Inches(footer_distance)
        except Exception:
            pass

        first_header = section.first_page_header
        first_header.is_linked_to_previous = False
        _clear_container_paragraphs(first_header)
        first_header_paragraph = first_header.add_paragraph()
        first_header_paragraph.add_run("")
        first_fmt = first_header_paragraph.paragraph_format
        first_fmt.space_before = Pt(0)
        first_fmt.space_after = Pt(0)

        header = section.header
        even_header = section.even_page_header

        header.is_linked_to_previous = False
        even_header.is_linked_to_previous = False
        _clear_container_paragraphs(header)
        _clear_container_paragraphs(even_header)

        if any_header:
            odd_paragraph = header.add_paragraph()
            odd_paragraph.alignment = odd_alignment
            odd_run = odd_paragraph.add_run(right_text or "")
            apply_run_format(odd_run, font_name=font_name, font_size_pt=font_size, make_italic=header_italic)
            if header_space_after:
                odd_paragraph.paragraph_format.space_after = Pt(header_space_after)

            even_paragraph = even_header.add_paragraph()
            even_paragraph.alignment = even_alignment
            even_run = even_paragraph.add_run(left_text or "")
            apply_run_format(even_run, font_name=font_name, font_size_pt=font_size, make_italic=header_italic)
            if header_space_after:
                even_paragraph.paragraph_format.space_after = Pt(header_space_after)
        else:
            header.add_paragraph("").add_run("")
            even_header.add_paragraph("").add_run("")

def add_table_of_contents(document: Document):
    """
    Inserta un campo de Tabla de Contenidos al principio del documento.
    """
    # AÃ±adir un pÃ¡rrafo donde irÃ¡ la TOC. Puede ser despuÃ©s del tÃ­tulo del libro.
    # AquÃ­ lo aÃ±adimos al principio para simplicidad.
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    # Crear el elemento fldChar para el inicio del campo complejo
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    # Crear la instrucciÃ³n del campo TOC
    instrText = OxmlElement('w:instrText')
    # '\o "1-3"' -> Incluye headings de nivel 1 a 3.
    # '\h' -> Crea hipervÃ­nculos.
    # '\z' -> Oculta la TOC en la vista web.
    # '\u' -> Usa los estilos de pÃ¡rrafo aplicados.
    instrText.text = r'TOC \o "1-3" \h \z \u'

    # Crear el elemento fldChar para el final del campo
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    # AÃ±adir los elementos al 'run'
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)

    # AÃ±adir un tÃ­tulo para la TOC
    toc_heading = document.add_paragraph('Ãndice', style='Heading 1')
    # Mover el tÃ­tulo antes del pÃ¡rrafo del campo TOC
    toc_heading._p.getparent().remove(toc_heading._p)
    paragraph._p.addprevious(toc_heading._p)

# -----------------------------
# API pÃºblica para importaciÃ³n
# -----------------------------

def _extract_front_matter(md_text: str) -> Dict[str, Any]:
    lines = md_text.splitlines()
    if not lines or lines[0].strip() != '---':
        return {}
    meta: Dict[str, Any] = {}
    for line in lines[1:]:
        if line.strip() == '---':
            break
        if ':' not in line:
            continue
        key, value = line.split(':', 1)
        key = key.strip()
        value = value.strip()
        if not key:
            continue
        try:
            meta[key] = json.loads(value)
        except Exception:
            meta[key] = value.strip('"')
    return meta

def _dict_to_block(item: Dict[str, Any]) -> Block:
    raw_type = (item.get("type") or "paragraph").lower()
    type_map = {"h1": "heading1", "h2": "heading2", "h3": "heading3", "p": "paragraph"}
    blk_type = type_map.get(raw_type, raw_type)
    content = item.get("text")
    if content is None:
        if blk_type in ("bullet", "numbered"):
            content = item.get("items", [])
        else:
            content = ""
    return Block(type=blk_type, content=content)

def block_to_dict(block: Block) -> Dict[str, Any]:
    """Convierte un objeto Block a la estructura de diccionario del pipeline."""
    type_map_reverse = {"heading1": "h1", "heading2": "h2", "heading3": "h3", "paragraph": "p"}
    doc_type = type_map_reverse.get(block.type, block.type)
    return {"type": doc_type, "text": block.content}


def convert_markdown_to_docx(
    md_path: Optional[str] = None,
    document: Optional[List[Any]] = None,
    output_path: Optional[str] = None,
    page_cfg: Optional[Dict[str, Any]] = None,
    style_cfgs: Optional[Dict[str, Dict[str, Any]]] = None
) -> str:
    """
    API principal para uso como mÃ³dulo.
    Puede recibir un `md_path` para leer y parsear, o una estructura `document` ya parseada.
    """
    if page_cfg is None:
        page_cfg = DEFAULT_PAGE.copy()
    if style_cfgs is None:
        style_cfgs = copy.deepcopy(DEFAULT_STYLESET)
    else:
        style_cfgs = copy.deepcopy(style_cfgs)

    blocks: List[Block]
    if document:
        blocks = []
        for item in document:
            if isinstance(item, Block):
                blocks.append(item)
            elif isinstance(item, dict):
                blocks.append(_dict_to_block(item))
            else:
                raise TypeError(f"Documento con tipo no soportado: {type(item)}")
        _apply_global_defaults(style_cfgs)
    elif md_path:
        with open(md_path, "r", encoding="utf-8") as f:
            md_text = f.read()
        blocks, _ = parse_markdown_blocks(md_text)

        front_meta = _extract_front_matter(md_text)
        title_meta = front_meta.get("title")
        author_meta = front_meta.get("author")
        _apply_global_defaults(style_cfgs, title_meta if isinstance(title_meta, str) else None, author_meta if isinstance(author_meta, str) else None)
    else:
        raise ValueError("Se debe proporcionar 'md_path' o 'document'.")

    doc = build_document(blocks, style_cfgs, page_cfg)
    configure_headers(doc, style_cfgs)
    add_page_numbers(doc, style_cfgs)

    if not output_path:
        if md_path:
            p = Path(md_path)
            output_path = str(p.with_name(p.stem + "_KDP.docx"))
        else:
            # Si no hay ruta de entrada, no se puede inferir la de salida
            raise ValueError("Se debe proporcionar 'output_path' si no se provee 'md_path'.")

    doc.save(output_path)
    return output_path

