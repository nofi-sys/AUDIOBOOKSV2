#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md2docx_gui.py — MVP
--------------------
Convierte un archivo Markdown (.md) a .docx con control de maquetado (tamaño de página, márgenes y estilos tipográficos por tipo de párrafo).
Pensado para KDP (Amazon). Incluye GUI con tkinter y arquitectura modular para importarlo como módulo desde otra app.

Características MVP:
- Detecta estilos básicos que aparecen en el .md (heading1/2/3, párrafo, cita, código, viñetas, numeración).
- Permite configurar tamaño de página, márgenes y, para cada estilo detectado, tipografía, tamaño, alineación, interlineado, espacio antes/después e indentación de primera línea.
- Exporta a .docx con texto en color negro.
- Diseño listo para reuso como módulo (función convert_markdown_to_docx).

Limitaciones MVP (para extender en futuras versiones):
- Listas anidadas limitadas (se maneja un único nivel).
- Negrita/itálica/subrayado/código en línea soportados, pero sin manejo de anidaciones complejas patológicas.
- Imágenes y tablas no incluidas en este MVP (se pueden añadir luego).
"""
import re
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import List, Dict, Tuple, Any, Optional, Set

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# -----------------------------
# Modelos de datos y defaults
# -----------------------------

STYLE_KEYS_ORDER = [
    "heading1",
    "heading2",
    "heading3",
    "paragraph",
    "blockquote",
    "codeblock",
    "bullet",
    "numbered",
]

DEFAULT_PAGE = {
    "preset": "KDP 5x8 in",
    "width_in": 5.0,
    "height_in": 8.0,
    "margin_top_in": 0.75,
    "margin_bottom_in": 0.75,
    "margin_left_in": 0.75,
    "margin_right_in": 0.75,
}

DEFAULT_STYLESET: Dict[str, Dict[str, Any]] = {
    "heading1": {
        "font_name": "Garamond",
        "font_size_pt": 16,
        "align": "left",
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "heading2": {
        "font_name": "Garamond",
        "font_size_pt": 14,
        "align": "left",
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "heading3": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "left",
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "paragraph": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "left",
        "line_spacing": 1.0,  # simple
        "space_before_pt": 0,
        "space_after_pt": 6,  # requisito: espacio después de párrafo por defecto
        "first_line_indent_in": 0.0,
    },
    "blockquote": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "left",
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "codeblock": {
        "font_name": "Courier New",
        "font_size_pt": 10,
        "align": "left",
        "line_spacing": 1.0,
        "space_before_pt": 6,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "bullet": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "left",
        "line_spacing": 1.0,
        "space_before_pt": 0,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
    "numbered": {
        "font_name": "Garamond",
        "font_size_pt": 12,
        "align": "left",
        "line_spacing": 1.0,
        "space_before_pt": 0,
        "space_after_pt": 6,
        "first_line_indent_in": 0.0,
    },
}

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

COMMON_FONTS = [
    "Garamond",
    "Times New Roman",
    "Georgia",
    "Palatino Linotype",
    "Cambria",
    "Book Antiqua",
    "Minion Pro",
    "Libre Baskerville",
    "Merriweather",
    "Courier New",
]

# -----------------------------
# Parser sencillo de Markdown
# -----------------------------

@dataclass
class Block:
    type: str  # 'heading1','heading2','heading3','paragraph','blockquote','codeblock','bullet','numbered','hr'
    content: Any  # str, list[str] (para listas), etc.

INLINE_TOKEN_RE = re.compile(
    r"(`[^`]+`)|(\*\*[^*]+\*\*)|(__[^_]+__)|(\*[^*]+\*)|(_[^_]+_)|(~~[^~]+~~)|(\+\+[^+]+\+\+)"
)

def tokenize_inline(text: str) -> List[Tuple[str, str]]:
    """
    Devuelve una lista de tokens inline (("text", texto) | ("bold","...") | ("italic","...") | ("code","...") | ("underline","...") | ("strike","...")).
    No maneja anidaciones patológicas. MVP suficiente.
    """
    tokens = []
    pos = 0
    for m in INLINE_TOKEN_RE.finditer(text):
        start, end = m.span()
        if start > pos:
            tokens.append(("text", text[pos:start]))
        group = m.group(0)
        if group.startswith("`"):
            tokens.append(("code", group[1:-1]))
        elif group.startswith("**") and group.endswith("**"):
            tokens.append(("bold", group[2:-2]))
        elif group.startswith("__") and group.endswith("__"):
            tokens.append(("bold", group[2:-2]))
        elif group.startswith("*") and group.endswith("*"):
            tokens.append(("italic", group[1:-1]))
        elif group.startswith("_") and group.endswith("_"):
            tokens.append(("italic", group[1:-1]))
        elif group.startswith("~~") and group.endswith("~~"):
            tokens.append(("strike", group[2:-2]))
        elif group.startswith("++") and group.endswith("++"):
            tokens.append(("underline", group[2:-2]))
        pos = end
    if pos < len(text):
        tokens.append(("text", text[pos:]))
    return tokens

FENCE_RE = re.compile(r"^```(\w+)?\s*$")
ATX_H_RE = re.compile(r"^(#{1,6})\s+(.*)$")
HR_RE = re.compile(r"^(\*\s*\*\s*\*|-{3,}|_{3,})\s*$")
BULLET_RE = re.compile(r"^([\-\*\+])\s+(.*)$")
ORDERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")

def parse_markdown_blocks(md_text: str) -> Tuple[List[Block], Set[str]]:
    """
    Parser de bloques muy simple para Markdown.
    Devuelve lista de Block y set de estilos detectados.
    """
    lines = md_text.splitlines()
    i = 0
    blocks: List[Block] = []
    used_styles: Set[str] = set()
    while i < len(lines):
        line = lines[i]

        # Código con fence ```
        m = FENCE_RE.match(line.strip())
        if m:
            lang = m.group(1) or ""
            i += 1
            buf = []
            while i < len(lines) and not FENCE_RE.match(lines[i].strip()):
                buf.append(lines[i])
                i += 1
            # Saltar la línea de cierre ``` si existe
            if i < len(lines) and FENCE_RE.match(lines[i].strip()):
                i += 1
            code_text = "\n".join(buf)
            blocks.append(Block("codeblock", code_text))
            used_styles.add("codeblock")
            continue

        # Regla horizontal
        if HR_RE.match(line.strip()):
            blocks.append(Block("hr", ""))
            i += 1
            continue

        # Encabezados ATX
        m = ATX_H_RE.match(line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                blocks.append(Block("heading1", text))
                used_styles.add("heading1")
            elif level == 2:
                blocks.append(Block("heading2", text))
                used_styles.add("heading2")
            elif level == 3:
                blocks.append(Block("heading3", text))
                used_styles.add("heading3")
            else:
                # niveles 4-6 caen como párrafo por MVP
                blocks.append(Block("paragraph", text))
                used_styles.add("paragraph")
            i += 1
            continue

        # Citas
        m = BLOCKQUOTE_RE.match(line)
        if m:
            buf = [m.group(1)]
            i += 1
            while i < len(lines):
                m2 = BLOCKQUOTE_RE.match(lines[i])
                if m2:
                    buf.append(m2.group(1))
                    i += 1
                else:
                    break
            text = "\n".join(buf).strip()
            blocks.append(Block("blockquote", text))
            used_styles.add("blockquote")
            continue

        # Listas con viñetas (un nivel)
        m = BULLET_RE.match(line)
        if m:
            items = [m.group(2)]
            i += 1
            while i < len(lines):
                m2 = BULLET_RE.match(lines[i])
                if m2:
                    items.append(m2.group(2))
                    i += 1
                else:
                    # permitir líneas en blanco entre ítems
                    if lines[i].strip() == "":
                        i += 1
                        continue
                    break
            blocks.append(Block("bullet", items))
            used_styles.add("bullet")
            continue

        # Listas numeradas (un nivel)
        m = ORDERED_RE.match(line)
        if m:
            items = [m.group(2)]
            i += 1
            while i < len(lines):
                m2 = ORDERED_RE.match(lines[i])
                if m2:
                    items.append(m2.group(2))
                    i += 1
                else:
                    if lines[i].strip() == "":
                        i += 1
                        continue
                    break
            blocks.append(Block("numbered", items))
            used_styles.add("numbered")
            continue

        # Párrafos (acumular hasta línea en blanco)
        if line.strip() == "":
            i += 1
            continue
        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip() != "":
            buf.append(lines[i])
            i += 1
        text = "\n".join(buf).strip()
        blocks.append(Block("paragraph", text))
        used_styles.add("paragraph")

    # Garantizar al menos 'paragraph' si no hubo nada
    if not used_styles:
        used_styles.add("paragraph")
    return blocks, used_styles

# -----------------------------
# Construcción del DOCX
# -----------------------------

def apply_paragraph_format(p, style_cfg: Dict[str, Any]):
    pf = p.paragraph_format
    pf.alignment = ALIGN_MAP.get(style_cfg.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)

    # Interlineado
    ls = style_cfg.get("line_spacing", 1.0)
    if isinstance(ls, (int, float)):
        # python-docx: usar line_spacing_rule + line_spacing si deseamos exacto/multiple
        # Como MVP, usaremos múltiplo aproximado: rule = MULTIPLE y valor = ls
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = ls
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Espacios antes/después
    before = style_cfg.get("space_before_pt", 0)
    after = style_cfg.get("space_after_pt", 0)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

    # Indentación primera línea
    first_indent_in = style_cfg.get("first_line_indent_in", 0.0)
    try:
        pf.first_line_indent = Inches(float(first_indent_in))
    except Exception:
        pf.first_line_indent = Inches(0.0)

def apply_run_format(run, make_bold=False, make_italic=False, make_underline=False, strike=False, font_name=None, font_size_pt=None):
    rfont = run.font
    rfont.bold = True if make_bold else None
    rfont.italic = True if make_italic else None
    rfont.underline = True if make_underline else None
    rfont.strike = True if strike else None
    if font_name:
        rfont.name = font_name
    if font_size_pt:
        rfont.size = Pt(font_size_pt)
    # Color negro
    rfont.color.rgb = RGBColor(0,0,0)

def write_inline_runs(p, text: str, base_font_name: str, base_font_size: int):
    tokens = tokenize_inline(text)
    for kind, val in tokens:
        run = p.add_run()
        if kind == "text":
            run.text = val
            apply_run_format(run, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "bold":
            run.text = val
            apply_run_format(run, make_bold=True, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "italic":
            run.text = val
            apply_run_format(run, make_italic=True, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "underline":
            run.text = val
            apply_run_format(run, make_underline=True, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "strike":
            run.text = val
            apply_run_format(run, strike=True, font_name=base_font_name, font_size_pt=base_font_size)
        elif kind == "code":
            # inline code: usar monoespaciada pequeña
            run.text = val
            apply_run_format(run, font_name="Courier New", font_size_pt=max(9, int(base_font_size*0.9)))

def setup_page(document: Document, page_cfg: Dict[str, Any]):
    for section in document.sections:
        section.page_width = Inches(float(page_cfg["width_in"]))
        section.page_height = Inches(float(page_cfg["height_in"]))
        section.top_margin = Inches(float(page_cfg["margin_top_in"]))
        section.bottom_margin = Inches(float(page_cfg["margin_bottom_in"]))
        section.left_margin = Inches(float(page_cfg["margin_left_in"]))
        section.right_margin = Inches(float(page_cfg["margin_right_in"]))

def apply_style_to_named_style(document: Document, style_name: str, cfg: Dict[str, Any]):
    """
    Ajusta parámetros básicos del estilo 'style_name' existente si existe;
    si no, intentará usar 'Normal' como base.
    """
    try:
        style = document.styles[style_name]
    except KeyError:
        style = document.styles['Normal']

    # Tipografía
    if hasattr(style, "font"):
        style.font.name = cfg.get("font_name", "Garamond")
        style.font.size = Pt(cfg.get("font_size_pt", 12))
        style.font.color.rgb = RGBColor(0,0,0)

    # Párrafo
    if hasattr(style, "paragraph_format"):
        pf = style.paragraph_format
        pf.alignment = ALIGN_MAP.get(cfg.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        pf.space_before = Pt(cfg.get("space_before_pt", 0))
        pf.space_after = Pt(cfg.get("space_after_pt", 6))
        pf.first_line_indent = Inches(cfg.get("first_line_indent_in", 0.0))
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = cfg.get("line_spacing", 1.0)

def build_document(blocks: List[Block], style_cfgs: Dict[str, Dict[str, Any]], page_cfg: Dict[str, Any]) -> Document:
    document = Document()

    # Página
    setup_page(document, page_cfg)

    # Acomodar estilos base en la plantilla
    # Heading 1/2/3, Normal, Quote, List Bullet, List Number, y crear uno para código si hace falta (usaremos Normal + Courier per párrafo)
    apply_style_to_named_style(document, 'Normal', style_cfgs.get("paragraph", DEFAULT_STYLESET["paragraph"]))
    apply_style_to_named_style(document, 'Heading 1', style_cfgs.get("heading1", DEFAULT_STYLESET["heading1"]))
    apply_style_to_named_style(document, 'Heading 2', style_cfgs.get("heading2", DEFAULT_STYLESET["heading2"]))
    apply_style_to_named_style(document, 'Heading 3', style_cfgs.get("heading3", DEFAULT_STYLESET["heading3"]))
    apply_style_to_named_style(document, 'Quote', style_cfgs.get("blockquote", DEFAULT_STYLESET["blockquote"]))
    apply_style_to_named_style(document, 'List Bullet', style_cfgs.get("bullet", DEFAULT_STYLESET["bullet"]))
    apply_style_to_named_style(document, 'List Number', style_cfgs.get("numbered", DEFAULT_STYLESET["numbered"]))

    for blk in blocks:
        if blk.type == "heading1":
            cfg = style_cfgs.get("heading1", DEFAULT_STYLESET["heading1"])
            p = document.add_paragraph(style='Heading 1')
            apply_paragraph_format(p, cfg)
            write_inline_runs(p, blk.content, cfg["font_name"], cfg["font_size_pt"])

        elif blk.type == "heading2":
            cfg = style_cfgs.get("heading2", DEFAULT_STYLESET["heading2"])
            p = document.add_paragraph(style='Heading 2')
            apply_paragraph_format(p, cfg)
            write_inline_runs(p, blk.content, cfg["font_name"], cfg["font_size_pt"])

        elif blk.type == "heading3":
            cfg = style_cfgs.get("heading3", DEFAULT_STYLESET["heading3"])
            p = document.add_paragraph(style='Heading 3')
            apply_paragraph_format(p, cfg)
            write_inline_runs(p, blk.content, cfg["font_name"], cfg["font_size_pt"])

        elif blk.type == "paragraph":
            cfg = style_cfgs.get("paragraph", DEFAULT_STYLESET["paragraph"])
            p = document.add_paragraph(style='Normal')
            apply_paragraph_format(p, cfg)
            write_inline_runs(p, blk.content, cfg["font_name"], cfg["font_size_pt"])

        elif blk.type == "blockquote":
            cfg = style_cfgs.get("blockquote", DEFAULT_STYLESET["blockquote"])
            p = document.add_paragraph(style='Quote')
            apply_paragraph_format(p, cfg)
            write_inline_runs(p, blk.content, cfg["font_name"], cfg["font_size_pt"])

        elif blk.type == "codeblock":
            cfg = style_cfgs.get("codeblock", DEFAULT_STYLESET["codeblock"])
            # Usar Normal pero forzar Courier y tamaño desde cfg en cada run
            p = document.add_paragraph(style='Normal')
            apply_paragraph_format(p, cfg)
            for line in blk.content.splitlines() or [""]:
                run = p.add_run(line)
                apply_run_format(run, font_name=cfg["font_name"], font_size_pt=cfg["font_size_pt"])
                # salto de línea manual para respetar líneas del bloque
                p.add_run("\n")

        elif blk.type == "bullet":
            cfg = style_cfgs.get("bullet", DEFAULT_STYLESET["bullet"])
            for item in blk.content:
                p = document.add_paragraph(style='List Bullet')
                apply_paragraph_format(p, cfg)
                write_inline_runs(p, item, cfg["font_name"], cfg["font_size_pt"])

        elif blk.type == "numbered":
            cfg = style_cfgs.get("numbered", DEFAULT_STYLESET["numbered"])
            for item in blk.content:
                p = document.add_paragraph(style='List Number')
                apply_paragraph_format(p, cfg)
                write_inline_runs(p, item, cfg["font_name"], cfg["font_size_pt"])

        elif blk.type == "hr":
            # Insertar separador como un párrafo vacío con espacio antes/después
            p = document.add_paragraph()
            run = p.add_run("• • •")
            apply_run_format(run, font_name='Garamond', font_size_pt=12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return document

# -----------------------------
# API pública para importación
# -----------------------------

def convert_markdown_to_docx(md_path: str, output_path: Optional[str] = None,
                             page_cfg: Optional[Dict[str, Any]] = None,
                             style_cfgs: Optional[Dict[str, Dict[str, Any]]] = None) -> str:
    """
    API principal para uso como módulo.
    """
    if page_cfg is None:
        page_cfg = DEFAULT_PAGE.copy()
    if style_cfgs is None:
        style_cfgs = DEFAULT_STYLESET.copy()

    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    blocks, _used = parse_markdown_blocks(md_text)
    doc = build_document(blocks, style_cfgs, page_cfg)

    if not output_path:
        p = Path(md_path)
        output_path = str(p.with_name(p.stem + "_KDP.docx"))
    doc.save(output_path)
    return output_path

# -----------------------------
# GUI con tkinter
# -----------------------------

class StylesDialog(tk.Toplevel):
    def __init__(self, master, used_styles: List[str], style_cfgs: Dict[str, Dict[str, Any]]):
        super().__init__(master)
        self.title("Configurar estilos tipográficos")
        self.style_cfgs = style_cfgs
        self.used_styles = used_styles
        self.vars: Dict[Tuple[str,str], tk.Variable] = {}
        self._build()

    def _build(self):
        container = ttk.Frame(self, padding=10)
        container.pack(fill="both", expand=True)

        info = ttk.Label(container, text="Configura tipografía y espaciamientos por estilo. (Color siempre negro en la exportación)")
        info.pack(anchor="w", pady=(0,10))

        canvas_container = ttk.Frame(container)
        canvas_container.pack(fill="both", expand=True)

        canvas = tk.Canvas(canvas_container, highlightthickness=0)
        scroll_y = ttk.Scrollbar(canvas_container, orient="vertical", command=canvas.yview)
        scroll_x = ttk.Scrollbar(container, orient="horizontal", command=canvas.xview)
        frame = ttk.Frame(canvas)
        frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        window_id = canvas.create_window((0, 0), window=frame, anchor="nw")

        canvas.configure(yscrollcommand=scroll_y.set, xscrollcommand=scroll_x.set)
        canvas.pack(side="left", fill="both", expand=True)
        scroll_y.pack(side="right", fill="y")
        scroll_x.pack(fill="x", pady=(6, 0))

        def _sync_inner_width(event, canvas=canvas, frame=frame, item_id=window_id):
            if frame.winfo_reqwidth() <= event.width:
                canvas.itemconfigure(item_id, width=event.width)
        canvas.bind("<Configure>", _sync_inner_width)

        for sk in self.used_styles:
            cfg = self.style_cfgs.get(sk, DEFAULT_STYLESET[sk])
            lf = ttk.LabelFrame(frame, text=f"Estilo: {sk}", padding=10)
            lf.pack(fill="x", expand=True, pady=6)

            # Font family
            ttk.Label(lf, text="Fuente:").grid(row=0, column=0, sticky="w")
            v_font = tk.StringVar(value=str(cfg.get("font_name","Garamond")))
            cb_font = ttk.Combobox(lf, textvariable=v_font, values=COMMON_FONTS, width=24)
            cb_font.grid(row=0, column=1, sticky="w", padx=6)

            # Font size
            ttk.Label(lf, text="Tamaño (pt):").grid(row=0, column=2, sticky="w")
            v_size = tk.DoubleVar(value=float(cfg.get("font_size_pt",12)))
            sp_size = ttk.Spinbox(lf, from_=8, to=48, increment=0.5, textvariable=v_size, width=6)
            sp_size.grid(row=0, column=3, sticky="w", padx=6)

            # Alignment
            ttk.Label(lf, text="Alineación:").grid(row=0, column=4, sticky="w")
            v_align = tk.StringVar(value=str(cfg.get("align","left")))
            cb_align = ttk.Combobox(lf, textvariable=v_align, values=["left","center","right","justify"], width=10)
            cb_align.grid(row=0, column=5, sticky="w", padx=6)

            # Line spacing
            ttk.Label(lf, text="Interlineado:").grid(row=1, column=0, sticky="w", pady=4)
            v_ls = tk.DoubleVar(value=float(cfg.get("line_spacing",1.0)))
            cb_ls = ttk.Combobox(lf, textvariable=v_ls, values=[1.0,1.15,1.5,2.0], width=6)
            cb_ls.grid(row=1, column=1, sticky="w", padx=6)

            # Space before
            ttk.Label(lf, text="Espacio antes (pt):").grid(row=1, column=2, sticky="w")
            v_before = tk.DoubleVar(value=float(cfg.get("space_before_pt",0)))
            sp_before = ttk.Spinbox(lf, from_=0, to=48, increment=1, textvariable=v_before, width=6)
            sp_before.grid(row=1, column=3, sticky="w", padx=6)

            # Space after
            ttk.Label(lf, text="Espacio después (pt):").grid(row=1, column=4, sticky="w")
            v_after = tk.DoubleVar(value=float(cfg.get("space_after_pt",6)))
            sp_after = ttk.Spinbox(lf, from_=0, to=48, increment=1, textvariable=v_after, width=6)
            sp_after.grid(row=1, column=5, sticky="w", padx=6)

            # First line indent
            ttk.Label(lf, text="Sangría primera línea (in):").grid(row=2, column=0, sticky="w", pady=4)
            v_indent = tk.DoubleVar(value=float(cfg.get("first_line_indent_in",0.0)))
            sp_indent = ttk.Spinbox(lf, from_=0.0, to=1.0, increment=0.05, textvariable=v_indent, width=6)
            sp_indent.grid(row=2, column=1, sticky="w", padx=6)

            # Guardar referencias
            self.vars[(sk,"font_name")] = v_font
            self.vars[(sk,"font_size_pt")] = v_size
            self.vars[(sk,"align")] = v_align
            self.vars[(sk,"line_spacing")] = v_ls
            self.vars[(sk,"space_before_pt")] = v_before
            self.vars[(sk,"space_after_pt")] = v_after
            self.vars[(sk,"first_line_indent_in")] = v_indent

        btns = ttk.Frame(container)
        btns.pack(fill="x", pady=(10,0))
        ttk.Button(btns, text="Aceptar", command=self.on_accept).pack(side="right")
        ttk.Button(btns, text="Cancelar", command=self.destroy).pack(side="right", padx=6)

        self.geometry("820x560")

    def on_accept(self):
        # Volcar valores al style_cfgs
        for (sk, key), var in self.vars.items():
            val = var.get()
            # cast numéricos
            if key in ("font_size_pt","line_spacing","space_before_pt","space_after_pt","first_line_indent_in"):
                try:
                    val = float(val)
                except Exception:
                    pass
            self.style_cfgs.setdefault(sk, {}).update({key: val})
        self.destroy()

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("MD → DOCX (KDP) — MVP")
        self.resizable(True, False)

        self.md_path: Optional[str] = None
        self.blocks: List[Block] = []
        self.used_styles: Set[str] = set()

        self.page_cfg = DEFAULT_PAGE.copy()
        self.style_cfgs = {k:v.copy() for k,v in DEFAULT_STYLESET.items()}

        self._build_ui()

    def _build_ui(self):
        root = ttk.Frame(self, padding=12)
        root.pack(fill="x", expand=True)

        # Selección de archivo
        file_row = ttk.Frame(root)
        file_row.pack(fill="x", pady=(0,8))
        ttk.Label(file_row, text="Archivo .md:").pack(side="left")
        self.var_path = tk.StringVar()
        entry = ttk.Entry(file_row, textvariable=self.var_path, width=70)
        entry.pack(side="left", padx=6, fill="x", expand=True)
        ttk.Button(file_row, text="Buscar…", command=self.choose_file).pack(side="left")

        # Página
        page_frame = ttk.LabelFrame(root, text="Página (KDP)", padding=10)
        page_frame.pack(fill="x", pady=(0,8))

        ttk.Label(page_frame, text="Preset:").grid(row=0, column=0, sticky="w")
        self.var_preset = tk.StringVar(value="KDP 5x8 in")
        cb_preset = ttk.Combobox(page_frame, textvariable=self.var_preset, values=[
            "KDP 5x8 in","KDP 5.25x8 in","KDP 6x9 in","A5","A4","Personalizado"
        ], width=15)
        cb_preset.grid(row=0, column=1, sticky="w", padx=6)
        cb_preset.bind("<<ComboboxSelected>>", self.on_preset_change)

        ttk.Label(page_frame, text="Ancho (in):").grid(row=0, column=2, sticky="w")
        self.var_w = tk.DoubleVar(value=self.page_cfg["width_in"])
        ttk.Spinbox(page_frame, from_=3.0, to=8.5, increment=0.05, textvariable=self.var_w, width=7).grid(row=0, column=3, sticky="w", padx=6)

        ttk.Label(page_frame, text="Alto (in):").grid(row=0, column=4, sticky="w")
        self.var_h = tk.DoubleVar(value=self.page_cfg["height_in"])
        ttk.Spinbox(page_frame, from_=5.0, to=11.7, increment=0.05, textvariable=self.var_h, width=7).grid(row=0, column=5, sticky="w", padx=6)

        ttk.Label(page_frame, text="Margen sup. (in):").grid(row=1, column=0, sticky="w", pady=(6,0))
        self.var_mt = tk.DoubleVar(value=self.page_cfg["margin_top_in"])
        ttk.Spinbox(page_frame, from_=0.25, to=1.5, increment=0.05, textvariable=self.var_mt, width=7).grid(row=1, column=1, sticky="w", padx=6, pady=(6,0))

        ttk.Label(page_frame, text="Margen inf. (in):").grid(row=1, column=2, sticky="w", pady=(6,0))
        self.var_mb = tk.DoubleVar(value=self.page_cfg["margin_bottom_in"])
        ttk.Spinbox(page_frame, from_=0.25, to=1.5, increment=0.05, textvariable=self.var_mb, width=7).grid(row=1, column=3, sticky="w", padx=6, pady=(6,0))

        ttk.Label(page_frame, text="Margen izq. (in):").grid(row=1, column=4, sticky="w", pady=(6,0))
        self.var_ml = tk.DoubleVar(value=self.page_cfg["margin_left_in"])
        ttk.Spinbox(page_frame, from_=0.25, to=1.5, increment=0.05, textvariable=self.var_ml, width=7).grid(row=1, column=5, sticky="w", padx=6, pady=(6,0))

        ttk.Label(page_frame, text="Margen der. (in):").grid(row=1, column=6, sticky="w", pady=(6,0))
        self.var_mr = tk.DoubleVar(value=self.page_cfg["margin_right_in"])
        ttk.Spinbox(page_frame, from_=0.25, to=1.5, increment=0.05, textvariable=self.var_mr, width=7).grid(row=1, column=7, sticky="w", padx=6, pady=(6,0))

        # Botones: Estilos + Convertir
        btn_row = ttk.Frame(root)
        btn_row.pack(fill="x", pady=(6,0))

        ttk.Button(btn_row, text="Configurar estilos tipográficos…", command=self.open_styles_dialog).pack(side="left")
        ttk.Button(btn_row, text="Convertir a DOCX", command=self.convert).pack(side="right")

        # Notas
        notes = ttk.Label(root, foreground="#444", wraplength=700,
            text="Nota: Este MVP soporta encabezados (#, ##, ###), párrafos, citas (>), listas simples (-, 1.), y bloques de código (``` ... ```). "
                 "El color de texto se fuerza a negro en todos los estilos. Las imágenes/tablas no están incluidas en esta versión.")
        notes.pack(fill="x", pady=(8,0))

    def on_preset_change(self, event=None):
        p = self.var_preset.get()
        presets = {
            "KDP 5x8 in": (5.0,8.0),
            "KDP 5.25x8 in": (5.25,8.0),
            "KDP 6x9 in": (6.0,9.0),
            "A5": (5.83, 8.27),
            "A4": (8.27, 11.69),
        }
        if p in presets:
            w,h = presets[p]
            self.var_w.set(w); self.var_h.set(h)

    def choose_file(self):
        path = filedialog.askopenfilename(title="Seleccionar Markdown", filetypes=[("Markdown","*.md"),("Texto","*.txt"),("Todos","*.*")])
        if not path:
            return
        self.md_path = path
        self.var_path.set(path)

        # Parse preliminar para detectar estilos usados
        try:
            with open(path, "r", encoding="utf-8") as f:
                txt = f.read()
            self.blocks, self.used_styles = parse_markdown_blocks(txt)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo leer/parsing: {e}")
            return

        # Reducir style_cfgs a los usados + defaults
        new_cfgs = {}
        for sk in STYLE_KEYS_ORDER:
            if sk in self.used_styles:
                new_cfgs[sk] = self.style_cfgs.get(sk, DEFAULT_STYLESET[sk]).copy()
        self.style_cfgs = new_cfgs

        messagebox.showinfo("Detectado", f"Se detectaron estilos: {', '.join(self.used_styles)}")

    def open_styles_dialog(self):
        if not self.used_styles:
            # si no se ha elegido archivo, mostrar defaults para párrafo
            self.used_styles = {"paragraph"}
            self.style_cfgs = {"paragraph": DEFAULT_STYLESET["paragraph"].copy()}
        dlg = StylesDialog(self, [sk for sk in STYLE_KEYS_ORDER if sk in self.used_styles], self.style_cfgs)
        self.wait_window(dlg)

    def sync_page_cfg(self):
        self.page_cfg["width_in"] = float(self.var_w.get())
        self.page_cfg["height_in"] = float(self.var_h.get())
        self.page_cfg["margin_top_in"] = float(self.var_mt.get())
        self.page_cfg["margin_bottom_in"] = float(self.var_mb.get())
        self.page_cfg["margin_left_in"] = float(self.var_ml.get())
        self.page_cfg["margin_right_in"] = float(self.var_mr.get())

    def convert(self):
        if not self.md_path:
            messagebox.showwarning("Falta archivo", "Seleccioná primero un archivo .md")
            return

        self.sync_page_cfg()

        try:
            with open(self.md_path, "r", encoding="utf-8") as f:
                md_text = f.read()
            blocks, _ = parse_markdown_blocks(md_text)
            doc = build_document(blocks, self.style_cfgs, self.page_cfg)
            out_path = str(Path(self.md_path).with_name(Path(self.md_path).stem + "_KDP.docx"))
            doc.save(out_path)
            messagebox.showinfo("Listo", f"Documento exportado:\n{out_path}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo convertir: {e}")

def main():
    # Si se usa como CLI rápida: md2docx_gui.py input.md [salida.docx]
    if len(sys.argv) >= 2 and sys.argv[1].lower().endswith(".md"):
        md_path = sys.argv[1]
        out_path = None
        if len(sys.argv) >= 3:
            out_path = sys.argv[2]
        result = convert_markdown_to_docx(md_path, out_path)
        print(f"Exportado: {result}")
        return

    # GUI
    app = App()
    app.mainloop()

if __name__ == "__main__":
    main()
